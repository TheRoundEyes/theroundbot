import os
import pandas as pd
import streamlit as st
from google.generativeai import GenerativeModel, GenerationConfig
import google.generativeai as genai
from dotenv import load_dotenv
import json
from docx import Document
import fitz  # PyMuPDF
import pytesseract
import pdfplumber
from langchain.text_splitter import CharacterTextSplitter
import concurrent.futures

st.set_page_config("The Round Bot 3.0 - Gemini",layout="wide")

load_dotenv()

api_key = st.text_input("Enter your API key", type="password")

genai.configure(api_key=api_key)

def chunk_text(text, chunk_size=2000, chunk_overlap=300, separator=" "):
    splitter = CharacterTextSplitter(
        separator=separator,
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap
    )
    return splitter.split_text(text)

def initialise_model(model_name="gemini-1.5-flash"):
    # Instantiate the model
    model = GenerativeModel(model_name="gemini-1.5-flash-001")

    # Define the generation configuration
    generation_config = GenerationConfig(
        temperature=0.0,  # Adjust temperature as needed
        response_mime_type="application/json"  # Ensures the response is in JSON format
    )
    return model

def extract_text_from_pdf(file):
    with pdfplumber.open(file) as pdf:
        text = [page.extract_text() for page in pdf.pages]
        return "\n\n".join(text)

# Function to extract text from DOC files
def extract_text_from_doc(file):
    doc = Document(file)
    full_text = []

    for para in doc.paragraphs:
        full_text.append(para.text)

    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text.append(cell.text)

    return '\n'.join(full_text)
def get_response(model, model_behaviour, extracted_text, prompt):
    all_dfs = []

    for chunk in chunk_text(extracted_text):
        try:
            response = model.generate_content([model_behaviour, chunk, prompt])
            return response.text
        except json.JSONDecodeError:
            st.error("Failed to parse JSON response.")
    #
    # if all_dfs:
    #     return pd.concat(all_dfs, axis=0)
    # else:

def process_file(uploaded_file):
    if uploaded_file.name.endswith(".pdf"):
        return extract_text_from_pdf(uploaded_file)
    elif uploaded_file.name.endswith(".docx"):
        return extract_text_from_doc(uploaded_file)
    return ""

model_behaviour = """
You are an expert in processing document structures and extracting data from them.
We will upload a DOC file, and you need to extract the relevant information based on the given prompt.
You will provide the output in table format. Text-based answers are not acceptable.
STRICTLY FOLLOW THE FORMAT BELOW. REMOVE ALL OUTPUT KEYS
Given a document, your task is to extract the text value of the following entities:
{
'Company Name':'',
'Client Name':'',
'Issue Date':'',
'Items':[
     {
          'Item Name':'',
          'Item Quantity':'',
          'Item Total':''
     }
    ],
}
If a Column is not present on the given columns, use the name of the column from the document. If you can't identify the value, leave it empty.
Format the JSON properly as it should be in a dictionary format.
Your output should be in JSON format and should only be called "table". 
Remove any other keys from the output.
"""


def compare_documents(doc1, doc2, model):

    model_behaviour = """
    You are an expert in analyzing and comparing documents.
    You are given two documents to compare and provide a detailed analysis and key differences of these documents.

    When you compare them you must follow the format below:
    Invoice Numbers: 'Difference'
    """

    # Wrap documents in a dictionary or list as expected by the model
    combined_input = f"{model_behaviour}\nDocument 1:\n{doc1}\n\nDocument 2:\n{doc2}\n"


    response = model.generate_content(combined_input)

    return response.text

model = initialise_model("gemini-1.5-flash")
st.header("The Round Bot 3.0 - Gemini")
# Read the prompt in text box
prompt = "Generate a table from the DOC file content"

with st.sidebar:
    # Interface to upload DOC file
    # Upload the first set of PDF or DOC files
    uploaded_files_set1 = st.file_uploader("Upload the first set of PDF or DOC files", type=["pdf", "docx"],
                                           accept_multiple_files=True)
    all_texts_set1 = []

    # Upload the second set of PDF or DOC files
    uploaded_files_set2 = st.file_uploader("Upload the second set of PDF or DOC files", type=["pdf", "docx"],
                                           accept_multiple_files=True)
    all_texts_set2 = []

col1, col2 = st.columns(2)

if uploaded_files_set1 and uploaded_files_set2:
    with st.spinner('Processing files...'):
        with concurrent.futures.ThreadPoolExecutor() as executor:
            results_set1 = executor.map(process_file, uploaded_files_set1)
            all_texts_set1 = list(results_set1)

            results_set2 = executor.map(process_file, uploaded_files_set2)
            all_texts_set2 = list(results_set2)

    # Combine all texts from both sets
    combined_text_set1 = ' '.join(all_texts_set1)
    combined_text_set2 = ' '.join(all_texts_set2)

    # Analyze both sets of texts using the LLM
    df_set1 = get_response(model,combined_text_set1, model_behaviour, prompt)
    df_set2 = get_response(model,combined_text_set2, model_behaviour, prompt)

    # Display the extracted data
    with col1:
        st.write("Extracted Data from Set 1:")
        with st.expander("Scroll to see more", expanded=True):
            st.write(df_set1)

    with col2:
        st.write("Extracted Data from Set 2:")
        with st.expander("Scroll to see more", expanded=True):
            st.write(df_set2)


    comparison = compare_documents(combined_text_set1, combined_text_set2,model)
    st.write(comparison)