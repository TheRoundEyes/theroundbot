import pdfplumber
from docx import Document

def extract_text_from_pdf(file):
    with pdfplumber.open(file) as pdf:
        text = [page.extract_text() for page in pdf.pages]
        return "\n\n".join(text)

# Function to extract text from DOC files
def extract_text_from_doc(file):
    doc = Document(file)
    full_text = []

    for para in doc.paragraphs:
        full_text.append(para.text)

    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text.append(cell.text)

    return '\n'.join(full_text)